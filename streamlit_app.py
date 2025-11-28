import streamlit as st
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime, date
import os
import math
import tempfile
import zipfile
from io import BytesIO
import base64
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    try:
        import PyMuPDF as fitz  # Alternativa de import
        PYMUPDF_AVAILABLE = True
    except ImportError:
        PYMUPDF_AVAILABLE = False
from PIL import Image

# Import the reference calculation functions
from EcoDataReferences import (
    checkLeftAtrium, checkAorta, checkSepto, checkLVWall, 
    checkLVChamber_d, checkLVChamber_s
)

# Page configuration
st.set_page_config(
    page_title="Vet Heart Care - Sistema de Laudos",
    page_icon="🩺",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.5rem;
        color: #2c3e50;
        margin-top: 2rem;
        margin-bottom: 1rem;
        border-bottom: 2px solid #3498db;
        padding-bottom: 0.5rem;
    }
    .info-box {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #3498db;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #28a745;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        color: #856404;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #ffc107;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Helper functions from Django views
def extract_weight(input_string):
    """Extract weight from string like '5.2 kg'"""
    if not input_string:
        return None
    
    kg_index = input_string.find("kg")
    if kg_index == -1:
        return None
    
    weight_str = input_string[:kg_index].strip()
    try:
        return int(float(weight_str))
    except ValueError:
        return None

def is_float(value):
    """Check if value can be converted to float"""
    try:
        float(value)
        return True
    except ValueError:
        return False

def string_to_float(string_value):
    """Convert string to float, return None if invalid"""
    if string_value and string_value.strip():
        try:
            return float(string_value)
        except ValueError:
            return None
    return None

def reference_results(echo_label, echo_value, echo_weight):
    """Calculate TAPSE reference results (simplified version)"""
    # This is a simplified version - in the original Django app it connects to PostgreSQL
    # For Streamlit, we'll use basic ranges based on weight
    if echo_weight < 5:
        min_val, max_val = 8, 12
    elif echo_weight < 15:
        min_val, max_val = 10, 14
    elif echo_weight < 30:
        min_val, max_val = 12, 16
    else:
        min_val, max_val = 14, 18
    
    if echo_value > min_val and echo_value < max_val:
        return "Normal"
    elif echo_value < min_val:
        return "Diminuido"
    elif echo_value > max_val:
        return "Aumentado"

def replace_text(text, data):
    """Replace placeholders in text with actual values"""
    for key, value in data.items():
        text = text.replace(key, str(value or ""))
    return text

def extract_images_from_pdf(pdf_path, progress_callback=None):
    """Extract all pages from PDF as images (renderiza cada página completa como imagem)"""
    images = []
    try:
        if not PYMUPDF_AVAILABLE:
            if progress_callback:
                progress_callback("❌ PyMuPDF não está disponível! Instale com: pip install PyMuPDF")
            return []
        
        # Usar PyMuPDF para renderizar páginas como imagens
        if progress_callback:
            progress_callback(f"📂 Abrindo PDF: {pdf_path}")
        
        pdf_document = fitz.open(pdf_path)
        num_pages = len(pdf_document)
        
        if progress_callback:
            progress_callback(f"📄 PDF tem {num_pages} página(s). Renderizando TODAS as páginas como imagens...")
            progress_callback(f"   (Incluindo páginas com texto, imagens, ou ambos)")
        
        for page_num in range(num_pages):
            try:
                if progress_callback:
                    progress_callback(f"🔄 Renderizando página {page_num + 1}/{num_pages}...")
                
                page = pdf_document[page_num]
                
                # Renderizar a página COMPLETA como imagem (incluindo tudo: texto, imagens, gráficos)
                # zoom=3.0 para alta qualidade (3x = 216 DPI, excelente para imagens)
                mat = fitz.Matrix(3.0, 3.0)  # 3x zoom para alta qualidade
                
                # Renderizar página completa - captura TUDO que está visível na página
                # alpha=False remove transparência (melhor para Word)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                # IMPORTANTE: Renderizar TODAS as páginas, mesmo que pareçam vazias
                # (algumas podem ter apenas imagens que são detectadas na renderização)
                
                # Salvar como PNG diretamente
                img_data = pix.tobytes("png")
                
                # Salvar imagem temporariamente
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                    tmp_img.write(img_data)
                    img_path = tmp_img.name
                    
                    images.append({
                        'path': img_path,
                        'page': page_num + 1,
                        'width': pix.width,
                        'height': pix.height,
                        'format': 'png'
                    })
                    
                    if progress_callback:
                        progress_callback(f"✅ Página {page_num + 1} renderizada e salva ({pix.width}x{pix.height}px)")
                
                pix = None  # Liberar memória
                
            except Exception as e:
                error_msg = f"❌ Erro ao renderizar página {page_num + 1}: {str(e)}"
                if progress_callback:
                    progress_callback(error_msg)
                import traceback
                if progress_callback:
                    progress_callback(f"Detalhes: {traceback.format_exc()}")
                continue
        
        pdf_document.close()
        
        if progress_callback:
            progress_callback(f"✅ Total de {len(images)} página(s) renderizada(s)!")
        
    except Exception as e:
        error_msg = f"❌ Erro ao processar PDF: {str(e)}"
        if progress_callback:
            progress_callback(error_msg)
        import traceback
        if progress_callback:
            progress_callback(f"Detalhes: {traceback.format_exc()}")
        return []
    
    return images

def process_pdf_data(table_data):
    """Process extracted PDF table data"""
    # Initialize variables
    owner_name = animal_name = breed = age = sex = weight = date_of_exam = operator_name = ""
    species = ""
    identification = ""
    description = ""
    report_date = ""
    
    # Cardiac measurements
    diastole_septo_iv = ventriculo_esq = diast_diad_ve = diast_par_post_ve = ""
    dia_sist_ve = frac_ejecao = massa_ve = indice_massa_ve = ""
    diad_interno_ve_diast_norm = dia_interno_ve_sist_norm = ""
    espessura_relat_parede = tricuspid = tapse = doppler = ""
    aorta = vel_pico_va = gp_max_va = veloc_pico_mit_onda_e = veloc_pico_mit_ond_a = ""
    grad_pico_mit_e = grad_pico_mitral_a = mtp_mitral = taxa_mitral_e_a = ""
    temp_desacel_onda_e_mitral = temp_relax_isovol_mitral = temp_relax_isovol_e_mitral = ""
    regurgitacao_mitral = vel_reg_mit = grad_reg_mitral = dpdt = ""
    regurgitacao_tric = vel_reg_tric = grad_reg_tric = arteria_pulmonar = ""
    vel_pico_pulmonar = gradiente_pico_pulmonar = tdi_mitral = onda_e_lateral = onda_a_lateral = ""
    razao_ee_lat = pressao_cap_pulmonar = vel_pico_mit_ond_e = ""
    aorta_atrio_esq = diam_aortico = diam_ae = dia_esfericidade = diad_diam_ve = ""
    diam_atrio_ao_esq = vmax_va = encurtamento_frac_VE = mapse = ""
    # Additional cardiac parameters
    frequencia_cardiaca = ritmo = ""

    # Process table data
    for row in table_data:
        for i in range(0, len(row), 1):
            if row[i] is not None and row[i] != '':
                key = row[i].strip()
                value = row[i + 1].strip() if i + 1 < len(row) and row[i + 1] is not None else None
                value2 = row[i + 2].strip() if i + 2 < len(row) and row[i + 2] is not None else None

                if key == 'Nome proprietário':
                    owner_name = value.replace("\n", " ").title() if value else ""
                elif key == 'Nome do animal':
                    animal_name = value if value else ""
                    if animal_name:
                        parts = animal_name.split(',')
                        if len(parts) > 1:
                            species = parts[1].strip()
                            animal_name = parts[0].strip()
                        else:
                            species = ""
                elif key == 'Raça':
                    breed = value if value else ""
                elif key == 'Idade':
                    age = value if value else ""
                elif key == 'Sexo':
                    sex = value if value else ""
                elif key == 'Peso':
                    weight = value if value else ""
                elif key == 'Identificação':
                    identification = value if value else ""
                elif key == 'Data do exame':
                    date_of_exam = value if value else ""
                elif key == 'Operador':
                    operator_name = value if value else ""
                elif key == 'Descrição do exame':
                    description = value if value else ""
                elif key == 'Data do relatório':
                    report_date = value if value else ""
                elif key == 'Diástole-Septo IV':
                    diastole_septo_iv = value if value else ""
                elif key == 'Ventrículo esq':
                    ventriculo_esq = value if value else ""
                elif key == 'Diást-diâmetro VE':
                    diast_diad_ve = value if value else ""
                elif key == 'Diástole Parede Post VE':
                    diast_par_post_ve = value if value else ""
                elif key == 'Diâmetro-síst VE':
                    dia_sist_ve = value if value else ""
                elif key == 'Fração Ejeção':
                    frac_ejecao = value if value else ""
                elif key == 'MAPSE':
                    mapse = value if value else ""
                elif key == 'Massa VE':
                    massa_ve = value if value else ""
                elif key == 'Índice Massa VE':
                    indice_massa_ve = value if value else ""
                elif key == 'Diâmetro interno VE diást norm':
                    diad_interno_ve_diast_norm = value if value else ""
                elif key == 'Diâmetro interno VE sist norm':
                    dia_interno_ve_sist_norm = value if value else ""
                elif key == 'Espessura relativa da parede':
                    espessura_relat_parede = value if value else ""
                elif key == 'Tricúspid':
                    tricuspid = value if value else ""
                elif key == 'TAPSE':
                    tapse = value if value else ""
                elif key == 'Doppler':
                    doppler = value if value else ""
                elif key == 'Aorta':
                    aorta = value if value else ""
                elif key == 'Vmáx VA':
                    vel_pico_va = value if value else ""
                elif key == 'GP máx VA':
                    gp_max_va = value if value else ""
                elif key == 'Vel Pico Mitral Onda E':
                    veloc_pico_mit_onda_e = value if value else ""
                elif key == 'Vel Pico Mit Ond A':
                    veloc_pico_mit_ond_a = value if value else ""
                elif key == 'Grad Pico Mit (E)':
                    grad_pico_mit_e = value if value else ""
                elif key == 'Gtad Pico Mitral (A)':
                    grad_pico_mitral_a = value if value else ""
                elif key == 'MTP Mitral':
                    mtp_mitral = value if value else ""
                elif key == 'Taxa Mitral E/A':
                    taxa_mitral_e_a = value if value else ""
                elif key == 'Temp Desacel Onda e Mitral':
                    temp_desacel_onda_e_mitral = value if value else ""
                elif key == 'Temp Relax Isovol Mitral':
                    temp_relax_isovol_mitral = value if value else ""
                elif key == 'Temp Relax Isovol E Mitral':
                    temp_relax_isovol_e_mitral = value if value else ""
                elif key == 'Regurgitação Mitral':
                    regurgitacao_mitral = value if value else ""
                elif key == 'Velocidade Reg Mit':
                    vel_reg_mit = value if value else ""
                elif key == 'Gradiente Reg Mitral':
                    grad_reg_mitral = value if value else ""
                elif key == 'dP/dt':
                    dpdt = value if value else ""
                elif key == 'Regurgitação Tric':
                    regurgitacao_tric = value if value else ""
                elif key == 'Vel Reg Tric':
                    vel_reg_tric = value if value else ""
                elif key == 'Grad Reg Tric':
                    grad_reg_tric = value if value else ""
                elif key == 'Artéria Pulmonar':
                    arteria_pulmonar = value if value else ""
                elif key == 'Vel pico Pulmonar':
                    vel_pico_pulmonar = value if value else ""
                elif key == 'Gradiente pico Pulmonar':
                    gradiente_pico_pulmonar = value if value else ""
                elif key == 'TDI Mitral':
                    tdi_mitral = value if value else ""
                elif key == 'Onda E\' Lateral':
                    onda_e_lateral = value if value else ""
                elif key == 'Onda A\' Lateral':
                    onda_a_lateral = value if value else ""
                elif key == 'Razão E\'/A\' Lat':
                    razao_ee_lat = value if value else ""
                elif key == 'Pressão capilar pulmonar':
                    pressao_cap_pulmonar = value if value else ""
                elif key == 'Aorta/átrio esq':
                    aorta_atrio_esq = value if value else ""
                elif key == 'Diâmetro Aórtico':
                    diam_aortico = value if value else ""
                elif key == 'Diâmetro AE':
                    diam_ae = value if value else ""
                elif key == 'Diâm Átrio/Ao Esq':
                    dia_esfericidade = value if value else ""
                elif key == 'Índice de esfericidade':
                    diad_diam_ve = value if value else ""
                elif key == 'Encurtamento Fracional VE':
                    encurtamento_frac_VE = value if value else ""
                elif key == 'Diâm Átrio/Ao Esq':
                    diam_atrio_ao_esq = value if value else ""
                elif key == 'Frequência Cardíaca':
                    frequencia_cardiaca = value if value else ""
                elif key == 'Ritmo':
                    ritmo = value if value else ""

    # Process species and gender
    if species:
        if species == 'FEL':
            species = 'Felina'
        elif species == 'CAN':
            species = 'Canina'
    
    if sex == 'F':
        sex = 'Fêmea'
    elif sex == 'M':
        sex = 'Macho'

    # Calculate reference results
    weight_num = extract_weight(weight) if weight else None
    tapse_resultado = ''
    diam_aortico_result = ''
    diam_ae_result = ''
    diastole_septo_iv_result = ''
    diast_par_post_ve_result = ''
    diast_diad_ve_result = ''
    dia_sist_ve_result = ''

    if weight_num and weight_num > 0:
        if is_float(tapse):
            tapse_fl = float(tapse)
            tapse_resultado = reference_results("tapse", tapse_fl, weight_num)
        
        if is_float(diam_aortico):
            diam_aortico_result = checkAorta(weight_num, float(diam_aortico))
        
        if is_float(diam_ae):
            diam_ae_result = checkLeftAtrium(weight_num, float(diam_ae))
        
        if is_float(diastole_septo_iv):
            diastole_septo_iv_result = checkSepto(weight_num, float(diastole_septo_iv))
        
        if is_float(diast_par_post_ve):
            diast_par_post_ve_result = checkLVWall(weight_num, float(diast_par_post_ve))
        
        if is_float(diast_diad_ve):
            diast_diad_ve_result = checkLVChamber_d(weight_num, float(diast_diad_ve))
        
        if is_float(dia_sist_ve):
            dia_sist_ve_result = checkLVChamber_s(weight_num, float(dia_sist_ve))

    return {
        'owner_name': owner_name,
        'animal_name': animal_name,
        'species': species,
        'breed': breed,
        'age': age,
        'sex': sex,
        'weight': weight,
        'date_of_exam': date_of_exam,
        'operator_name': operator_name,
        'identification': identification,
        'description': description,
        'report_date': report_date,
        'diastole_septo_iv': diastole_septo_iv,
        'diast_par_post_ve': diast_par_post_ve,
        'diast_diad_ve': diast_diad_ve,
        'dia_sist_ve': dia_sist_ve,
        'encurtamento_frac_VE': encurtamento_frac_VE,
        'frac_ejecao': frac_ejecao,
        'diad_interno_ve_diast_norm': diad_interno_ve_diast_norm,
        'diam_aortico': diam_aortico,
        'diam_ae': diam_ae,
        'diam_atrio_ao_esq': diam_atrio_ao_esq,
        'vel_pico_va': vel_pico_va,
        'gp_max_va': gp_max_va,
        'vel_pico_pulmonar': vel_pico_pulmonar,
        'gradiente_pico_pulmonar': gradiente_pico_pulmonar,
        'veloc_pico_mit_onda_e': veloc_pico_mit_onda_e,
        'veloc_pico_mit_ond_a': veloc_pico_mit_ond_a,
        'taxa_mitral_e_a': taxa_mitral_e_a,
        'temp_desacel_onda_e_mitral': temp_desacel_onda_e_mitral,
        'temp_relax_isovol_mitral': temp_relax_isovol_mitral,
        'temp_relax_isovol_e_mitral': temp_relax_isovol_e_mitral,
        'onda_e_lateral': onda_e_lateral,
        'onda_a_lateral': onda_a_lateral,
        'razao_ee_lat': razao_ee_lat,
        'tapse': tapse,
        'mapse': mapse,
        'frequencia_cardiaca': frequencia_cardiaca,
        'ritmo': ritmo,
        'ventriculo_esq': ventriculo_esq,
        'massa_ve': massa_ve,
        'indice_massa_ve': indice_massa_ve,
        'dia_interno_ve_sist_norm': dia_interno_ve_sist_norm,
        'espessura_relat_parede': espessura_relat_parede,
        'tricuspid': tricuspid,
        'doppler': doppler,
        'aorta': aorta,
        'grad_pico_mit_e': grad_pico_mit_e,
        'grad_pico_mitral_a': grad_pico_mitral_a,
        'mtp_mitral': mtp_mitral,
        'regurgitacao_mitral': regurgitacao_mitral,
        'vel_reg_mit': vel_reg_mit,
        'grad_reg_mitral': grad_reg_mitral,
        'dpdt': dpdt,
        'regurgitacao_tric': regurgitacao_tric,
        'vel_reg_tric': vel_reg_tric,
        'grad_reg_tric': grad_reg_tric,
        'arteria_pulmonar': arteria_pulmonar,
        'tdi_mitral': tdi_mitral,
        'pressao_cap_pulmonar': pressao_cap_pulmonar,
        'aorta_atrio_esq': aorta_atrio_esq,
        'dia_esfericidade': dia_esfericidade,
        'diad_diam_ve': diad_diam_ve,
        'tapse_resultado': tapse_resultado,
        'diam_aortico_result': diam_aortico_result,
        'diam_ae_result': diam_ae_result,
        'diastole_septo_iv_result': diastole_septo_iv_result,
        'diast_par_post_ve_result': diast_par_post_ve_result,
        'diast_diad_ve_result': diast_diad_ve_result,
        'dia_sist_ve_result': dia_sist_ve_result,
    }

def create_word_document(data, template_path, output_path, images=None):
    """Create Word document from template and optionally add images from PDF"""
    try:
        doc = Document(template_path)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = replace_text(run.text, data)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace(key, str(value or ""))
        
        # Add user-uploaded images in grids of 6 per page
        if images and len(images) > 0:
            st.write(f"🔍 Inserindo {len(images)} imagem(ns) em blocos de até 6 por página...")
            try:
                # Adicionar quebra de página antes das imagens
                doc.add_page_break()
                
                # Adicionar título para a seção de imagens
                title_para = doc.add_paragraph()
                try:
                    title_para.alignment = 1  # Centralizado
                except:
                    pass
                title_run = title_para.add_run('Imagens Anexas')
                try:
                    title_run.font.size = Pt(16)
                    title_run.font.bold = True
                except:
                    pass
                
                doc.add_paragraph()
                
                def chunk_list(seq, size):
                    for i in range(0, len(seq), size):
                        yield seq[i:i+size]
                
                images_added = 0
                chunks = list(chunk_list(images, 6))
                for chunk_index, chunk in enumerate(chunks):
                    # 2 colunas x 3 linhas para imagens maiores
                    table = doc.add_table(rows=3, cols=2)
                    try:
                        table.alignment = 1
                    except:
                        pass
                    
                    for idx, img_info in enumerate(chunk):
                        cell = table.cell(idx // 2, idx % 2)
                        paragraph = cell.paragraphs[0]
                        try:
                            paragraph.alignment = 1
                        except:
                            pass
                        run = paragraph.add_run()
                        
                        img_path = img_info.get('path')
                        if not img_path or not os.path.exists(img_path):
                            st.warning(f"Imagem não encontrada: {img_path}")
                            continue
                        
                        width = img_info.get('width')
                        height = img_info.get('height')
                        max_width_inches = 3.2  # 2 colunas
                        max_height_inches = 3.5  # 3 linhas
                        # Escala respeitando proporção e limitando altura/largura para não quebrar página
                        if width and height:
                            ratio_w = max_width_inches / float(width)
                            ratio_h = max_height_inches / float(height)
                            scale = min(ratio_w, ratio_h)
                            width_inches = max_width_inches if scale >= 1 else max_width_inches * scale
                        else:
                            width_inches = max_width_inches
                        
                        try:
                            with open(img_path, "rb") as f:
                                img_bytes = BytesIO(f.read())
                                img_bytes.seek(0)
                            run.add_picture(img_bytes, width=Inches(width_inches))
                        except Exception as e:
                            st.error(f"Erro ao inserir imagem {img_info.get('name') or idx + 1}: {e}")
                            continue
                        
                        caption = cell.add_paragraph(img_info.get('name') or f"Imagem {images_added + 1}")
                        try:
                            caption.alignment = 1
                            caption.runs[0].font.size = Pt(9)
                        except:
                            pass
                        
                        images_added += 1
                    
                    if chunk_index < len(chunks) - 1:
                        doc.add_page_break()
                
                if images_added == 0:
                    st.warning("⚠️ Nenhuma imagem foi adicionada ao documento.")
                
                # Salvar documento com imagens
                doc.save(output_path)
            
            except Exception as e:
                st.error(f"Erro ao processar imagens: {str(e)}")
                import traceback
                st.error(f"Detalhes: {traceback.format_exc()}")
                doc.save(output_path)
        else:
            # Sem imagens, apenas salvar o documento normalmente
            doc.save(output_path)
        
        return True
    except Exception as e:
        st.error(f"Erro ao criar documento: {str(e)}")
        return False

def main():
    # Main header
    st.markdown('<h1 class="main-header">🩺 Vet Heart Care - Sistema de Laudos</h1>', unsafe_allow_html=True)
    
    # Initialize session state
    if 'step' not in st.session_state:
        st.session_state.step = 1
    if 'extracted_data' not in st.session_state:
        st.session_state.extracted_data = {}
    if 'pdf_file' not in st.session_state:
        st.session_state.pdf_file = None
    if 'pdf_bytes' not in st.session_state:
        st.session_state.pdf_bytes = None

    # Sidebar navigation
    st.sidebar.title("Navegação")
    steps = ["1. Upload do PDF", "2. Confirmação de Dados", "3. Geração do Laudo"]
    
    for i, step in enumerate(steps, 1):
        if i <= st.session_state.step:
            st.sidebar.success(step)
        else:
            st.sidebar.info(step)

    # Step 1: PDF Upload
    if st.session_state.step == 1:
        st.markdown('<h2 class="section-header">📄 Upload do Relatório Ecocardiográfico</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        <div class="info-box">
        <strong>Instruções:</strong><br>
        1. Faça upload do arquivo PDF do relatório ecocardiográfico<br>
        2. O sistema irá extrair automaticamente os dados do relatório<br>
        3. Você poderá revisar e confirmar os dados na próxima etapa
        </div>
        """, unsafe_allow_html=True)
        
        uploaded_file = st.file_uploader(
            "Selecione o arquivo PDF do relatório ecocardiográfico",
            type=['pdf'],
            help="Apenas arquivos PDF são aceitos"
        )
        
        if uploaded_file is not None:
            # Guardar bytes do PDF no session_state para usar depois
            pdf_bytes = uploaded_file.getvalue()
            st.session_state.pdf_bytes = pdf_bytes
            
            # Salvar arquivo temporário para extração de dados
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(pdf_bytes)
                tmp_file_path = tmp_file.name
            
            st.session_state.pdf_file = tmp_file_path
            
            # Extract data from PDF
            with st.spinner("Extraindo dados do PDF..."):
                try:
                    table_data = []
                    with pdfplumber.open(tmp_file_path) as pdf:
                        for page in pdf.pages:
                            tables = page.extract_tables()
                            for table in tables:
                                if table:
                                    table_data.extend(table)
                    
                    if table_data:
                        # Process the extracted data
                        extracted_data = process_pdf_data(table_data)
                        st.session_state.extracted_data = extracted_data
                        
                        st.success("✅ Dados extraídos com sucesso!")
                        st.info(f"Dados encontrados para: {extracted_data.get('animal_name', 'Animal não identificado')}")
                        
                        # Show preview of extracted data
                        with st.expander("🔍 Prévia dos Dados Extraídos", expanded=False):
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.write("**Informações do Animal:**")
                                st.write(f"- Nome: {extracted_data.get('animal_name', 'N/A')}")
                                st.write(f"- Espécie: {extracted_data.get('species', 'N/A')}")
                                st.write(f"- Raça: {extracted_data.get('breed', 'N/A')}")
                                st.write(f"- Sexo: {extracted_data.get('sex', 'N/A')}")
                                st.write(f"- Peso: {extracted_data.get('weight', 'N/A')}")
                                st.write(f"- Data do Exame: {extracted_data.get('date_of_exam', 'N/A')}")
                            
                            with col2:
                                st.write("**Informações do Proprietário:**")
                                st.write(f"- Nome: {extracted_data.get('owner_name', 'N/A')}")
                                st.write(f"- Operador: {extracted_data.get('operator_name', 'N/A')}")
                                st.write(f"- Identificação: {extracted_data.get('identification', 'N/A')}")
                        
                        if st.button("➡️ Continuar para Confirmação", type="primary"):
                            st.session_state.step = 2
                            st.rerun()
                    else:
                        st.error("❌ Não foi possível extrair dados do PDF. Verifique se o arquivo contém tabelas com dados.")
                
                except Exception as e:
                    st.error(f"❌ Erro ao processar o PDF: {str(e)}")
                
                finally:
                    # Clean up temporary file
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)

    # Step 2: Data Confirmation
    elif st.session_state.step == 2:
        st.markdown('<h2 class="section-header">✅ Confirmação de Dados</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        <div class="info-box">
        <strong>Revisão e Edição:</strong><br>
        Revise os dados extraídos do PDF e faça as correções necessárias antes de gerar o laudo.
        </div>
        """, unsafe_allow_html=True)
        
        # Back button
        if st.button("⬅️ Voltar ao Upload"):
            st.session_state.step = 1
            st.rerun()
        
        extracted_data = st.session_state.extracted_data
        
        # Create form for data confirmation
        with st.form("data_confirmation_form"):
            st.subheader("📋 Informações do Animal e Proprietário")
            
            col1, col2 = st.columns(2)
            
            with col1:
                owner_name = st.text_input(
                    "Nome do Proprietário",
                    value=extracted_data.get('owner_name', ''),
                    help="Nome completo do proprietário do animal"
                )
                
                animal_name = st.text_input(
                    "Nome do Animal",
                    value=extracted_data.get('animal_name', ''),
                    help="Nome do animal"
                )
                
                species = st.selectbox(
                    "Espécie",
                    options=["", "Canina", "Felina"],
                    index=["", "Canina", "Felina"].index(extracted_data.get('species', '')) if extracted_data.get('species', '') in ["", "Canina", "Felina"] else 0,
                    help="Espécie do animal"
                )
                
                breed = st.text_input(
                    "Raça",
                    value=extracted_data.get('breed', ''),
                    help="Raça do animal"
                )
                
                age = st.text_input(
                    "Idade",
                    value=extracted_data.get('age', ''),
                    help="Idade do animal"
                )
                
                sex = st.selectbox(
                    "Sexo",
                    options=["", "Macho", "Fêmea"],
                    index=["", "Macho", "Fêmea"].index(extracted_data.get('sex', '')) if extracted_data.get('sex', '') in ["", "Macho", "Fêmea"] else 0,
                    help="Sexo do animal"
                )
            
            with col2:
                weight = st.text_input(
                    "Peso",
                    value=extracted_data.get('weight', ''),
                    help="Peso do animal (ex: 5.2 kg)"
                )
                
                date_of_exam = st.text_input(
                    "Data do Exame",
                    value=extracted_data.get('date_of_exam', ''),
                    help="Data do exame (formato: DD/MM/AAAA)"
                )
                
                operator_name = st.text_input(
                    "Operador",
                    value=extracted_data.get('operator_name', ''),
                    help="Nome do operador do exame"
                )
                
                identification = st.text_input(
                    "Número da Ficha",
                    value=extracted_data.get('identification', ''),
                    help="Número de identificação da ficha"
                )
                
                email = st.text_input(
                    "Email do Proprietário",
                    value="",
                    help="Email para contato (opcional)"
                )
                
                phone = st.text_input(
                    "Telefone do Proprietário",
                    value="",
                    help="Telefone para contato (opcional)"
                )
                
                description = st.text_area(
                    "Descrição do Exame",
                    value=extracted_data.get('description', ''),
                    help="Descrição detalhada do exame ecocardiográfico",
                    height=80
                )
                
                report_date = st.text_input(
                    "Data do Relatório",
                    value=extracted_data.get('report_date', ''),
                    help="Data de geração do relatório"
                )
            
            st.subheader("📊 Dados Ecocardiográficos")
            
            # Frequência cardíaca e ritmo
            col1, col2 = st.columns(2)
            with col1:
                frequencia_cardiaca = st.text_input(
                    "Frequência Cardíaca",
                    value=extracted_data.get('frequencia_cardiaca', ''),
                    help="Frequência cardíaca em bpm"
                )
            with col2:
                ritmo = st.selectbox(
                    "Ritmo",
                    options=["", "Ritmo Sinusal", "Arritmia Sinusal", "Extrasístoles", "Fibrilação Atrial", "Outro"],
                    index=["", "Ritmo Sinusal", "Arritmia Sinusal", "Extrasístoles", "Fibrilação Atrial", "Outro"].index(extracted_data.get('ritmo', '')) if extracted_data.get('ritmo', '') in ["", "Ritmo Sinusal", "Arritmia Sinusal", "Extrasístoles", "Fibrilação Atrial", "Outro"] else 0,
                    help="Tipo de ritmo cardíaco"
                )
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.write("**Medidas Básicas:**")
                diastole_septo_iv = st.text_input("Diástole-Septo IV", value=extracted_data.get('diastole_septo_iv', ''))
                diast_par_post_ve = st.text_input("Diástole Parede Post VE", value=extracted_data.get('diast_par_post_ve', ''))
                diast_diad_ve = st.text_input("Diást-diâmetro VE", value=extracted_data.get('diast_diad_ve', ''))
                dia_sist_ve = st.text_input("Diâmetro-síst VE", value=extracted_data.get('dia_sist_ve', ''))
                encurtamento_frac_VE = st.text_input("Encurtamento Fracional VE", value=extracted_data.get('encurtamento_frac_VE', ''))
                frac_ejecao = st.text_input("Fração Ejeção", value=extracted_data.get('frac_ejecao', ''))
            
            with col2:
                st.write("**Medidas Adicionais:**")
                diad_interno_ve_diast_norm = st.text_input("Diâmetro interno VE diást norm", value=extracted_data.get('diad_interno_ve_diast_norm', ''))
                diam_aortico = st.text_input("Diâmetro Aórtico", value=extracted_data.get('diam_aortico', ''))
                diam_ae = st.text_input("Diâmetro AE", value=extracted_data.get('diam_ae', ''))
                diam_atrio_ao_esq = st.text_input("Diâm Átrio/Ao Esq", value=extracted_data.get('diam_atrio_ao_esq', ''))
                vel_pico_va = st.text_input("Vmáx VA", value=extracted_data.get('vel_pico_va', ''))
                gp_max_va = st.text_input("GP máx VA", value=extracted_data.get('gp_max_va', ''))
            
            with col3:
                st.write("**Doppler e TAPSE:**")
                vel_pico_pulmonar = st.text_input("Vel pico Pulmonar", value=extracted_data.get('vel_pico_pulmonar', ''))
                gradiente_pico_pulmonar = st.text_input("Gradiente pico Pulmonar", value=extracted_data.get('gradiente_pico_pulmonar', ''))
                veloc_pico_mit_onda_e = st.text_input("Vel Pico Mitral Onda E", value=extracted_data.get('veloc_pico_mit_onda_e', ''))
                veloc_pico_mit_ond_a = st.text_input("Vel Pico Mit Ond A", value=extracted_data.get('veloc_pico_mit_ond_a', ''))
                taxa_mitral_e_a = st.text_input("Taxa Mitral E/A", value=extracted_data.get('taxa_mitral_e_a', ''))
                tapse = st.text_input("TAPSE", value=extracted_data.get('tapse', ''))
            
            # Campos adicionais em seção expandível
            with st.expander("🔬 Medidas Adicionais Ecocardiográficas", expanded=False):
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.write("**Medidas Estruturais:**")
                    massa_ve = st.text_input("Massa VE", value=extracted_data.get('massa_ve', ''))
                    indice_massa_ve = st.text_input("Índice Massa VE", value=extracted_data.get('indice_massa_ve', ''))
                    dia_interno_ve_sist_norm = st.text_input("Diâmetro interno VE sist norm", value=extracted_data.get('dia_interno_ve_sist_norm', ''))
                    espessura_relat_parede = st.text_input("Espessura relativa da parede", value=extracted_data.get('espessura_relat_parede', ''))
                    ventriculo_esq = st.text_input("Ventrículo esq", value=extracted_data.get('ventriculo_esq', ''))
                
                with col2:
                    st.write("**Valvas e Doppler:**")
                    tricuspid = st.text_input("Tricúspid", value=extracted_data.get('tricuspid', ''))
                    doppler = st.text_input("Doppler", value=extracted_data.get('doppler', ''))
                    aorta = st.text_input("Aorta", value=extracted_data.get('aorta', ''))
                    arteria_pulmonar = st.text_input("Artéria Pulmonar", value=extracted_data.get('arteria_pulmonar', ''))
                    tdi_mitral = st.text_input("TDI Mitral", value=extracted_data.get('tdi_mitral', ''))
                
                with col3:
                    st.write("**Gradientes e Pressões:**")
                    grad_pico_mit_e = st.text_input("Grad Pico Mit (E)", value=extracted_data.get('grad_pico_mit_e', ''))
                    grad_pico_mitral_a = st.text_input("Grad Pico Mitral (A)", value=extracted_data.get('grad_pico_mitral_a', ''))
                    mtp_mitral = st.text_input("MTP Mitral", value=extracted_data.get('mtp_mitral', ''))
                    dpdt = st.text_input("dP/dt", value=extracted_data.get('dpdt', ''))
                    pressao_cap_pulmonar = st.text_input("Pressão capilar pulmonar", value=extracted_data.get('pressao_cap_pulmonar', ''))
            
            # Regurgitações em seção separada
            with st.expander("🔄 Regurgitações e Fluxos", expanded=False):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write("**Regurgitação Mitral:**")
                    regurgitacao_mitral = st.text_input("Regurgitação Mitral", value=extracted_data.get('regurgitacao_mitral', ''))
                    vel_reg_mit = st.text_input("Velocidade Reg Mit", value=extracted_data.get('vel_reg_mit', ''))
                    grad_reg_mitral = st.text_input("Gradiente Reg Mitral", value=extracted_data.get('grad_reg_mitral', ''))
                
                with col2:
                    st.write("**Regurgitação Tricúspide:**")
                    regurgitacao_tric = st.text_input("Regurgitação Tric", value=extracted_data.get('regurgitacao_tric', ''))
                    vel_reg_tric = st.text_input("Vel Reg Tric", value=extracted_data.get('vel_reg_tric', ''))
                    grad_reg_tric = st.text_input("Grad Reg Tric", value=extracted_data.get('grad_reg_tric', ''))
            
            # Campos adicionais
            with st.expander("📏 Medidas Adicionais", expanded=False):
                col1, col2 = st.columns(2)
                
                with col1:
                    aorta_atrio_esq = st.text_input("Aorta/átrio esq", value=extracted_data.get('aorta_atrio_esq', ''))
                    dia_esfericidade = st.text_input("Diâm Átrio/Ao Esq", value=extracted_data.get('dia_esfericidade', ''))
                    diad_diam_ve = st.text_input("Índice de esfericidade", value=extracted_data.get('diad_diam_ve', ''))
                
                with col2:
                    temp_desacel_onda_e_mitral = st.text_input("Temp Desacel Onda e Mitral", value=extracted_data.get('temp_desacel_onda_e_mitral', ''))
                    temp_relax_isovol_mitral = st.text_input("Temp Relax Isovol Mitral", value=extracted_data.get('temp_relax_isovol_mitral', ''))
                    temp_relax_isovol_e_mitral = st.text_input("Temp Relax Isovol E Mitral", value=extracted_data.get('temp_relax_isovol_e_mitral', ''))
            
            # Seção de resultados de referência calculados
            st.subheader("📈 Resultados de Referência Calculados")
            
            # Mostrar resultados se disponíveis
            if any([
                extracted_data.get('tapse_resultado'),
                extracted_data.get('diam_aortico_result'),
                extracted_data.get('diam_ae_result'),
                extracted_data.get('diastole_septo_iv_result'),
                extracted_data.get('diast_par_post_ve_result'),
                extracted_data.get('diast_diad_ve_result'),
                extracted_data.get('dia_sist_ve_result')
            ]):
                st.markdown("""
                <div class="info-box">
                <strong>Referências Médicas Calculadas:</strong><br>
                Os valores abaixo são calculados automaticamente baseados no peso do animal e nas medidas ecocardiográficas.
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if extracted_data.get('tapse_resultado'):
                        st.metric("TAPSE", extracted_data.get('tapse_resultado'), 
                                help="Referência baseada no peso do animal")
                    if extracted_data.get('diam_aortico_result'):
                        st.metric("Diâmetro Aórtico", extracted_data.get('diam_aortico_result'),
                                help="Referência baseada no peso do animal")
                
                with col2:
                    if extracted_data.get('diam_ae_result'):
                        st.metric("Diâmetro AE", extracted_data.get('diam_ae_result'),
                                help="Referência baseada no peso do animal")
                    if extracted_data.get('diastole_septo_iv_result'):
                        st.metric("Diástole-Septo IV", extracted_data.get('diastole_septo_iv_result'),
                                help="Referência baseada no peso do animal")
                
                with col3:
                    if extracted_data.get('diast_par_post_ve_result'):
                        st.metric("Diástole Parede Post VE", extracted_data.get('diast_par_post_ve_result'),
                                help="Referência baseada no peso do animal")
                    if extracted_data.get('diast_diad_ve_result'):
                        st.metric("Diást-diâmetro VE", extracted_data.get('diast_diad_ve_result'),
                                help="Referência baseada no peso do animal")
                    if extracted_data.get('dia_sist_ve_result'):
                        st.metric("Diâmetro-síst VE", extracted_data.get('dia_sist_ve_result'),
                                help="Referência baseada no peso do animal")
            else:
                st.info("ℹ️ Os resultados de referência serão calculados após confirmar os dados e peso do animal.")
            
            st.subheader("📝 Observações e Conclusões")
            
            col1, col2 = st.columns(2)
            
            with col1:
                notes = st.text_area(
                    "Observações",
                    value="",
                    help="Observações adicionais sobre o exame",
                    height=100
                )
            
            with col2:
                conclusion = st.text_area(
                    "Conclusões",
                    value="",
                    help="Conclusões do exame ecocardiográfico",
                    height=100
                )
            
            # Submit button
            submitted = st.form_submit_button("✅ Confirmar Dados e Gerar Laudo", type="primary")
            
            if submitted:
                # Update session state with form data
                st.session_state.extracted_data.update({
                    'owner_name': owner_name,
                    'animal_name': animal_name,
                    'species': species,
                    'breed': breed,
                    'age': age,
                    'sex': sex,
                    'weight': weight,
                    'date_of_exam': date_of_exam,
                    'operator_name': operator_name,
                    'identification': identification,
                    'email': email,
                    'phone': phone,
                    'description': description,
                    'report_date': report_date,
                    'frequencia_cardiaca': frequencia_cardiaca,
                    'ritmo': ritmo,
                    'diastole_septo_iv': diastole_septo_iv,
                    'diast_par_post_ve': diast_par_post_ve,
                    'diast_diad_ve': diast_diad_ve,
                    'dia_sist_ve': dia_sist_ve,
                    'encurtamento_frac_VE': encurtamento_frac_VE,
                    'frac_ejecao': frac_ejecao,
                    'diad_interno_ve_diast_norm': diad_interno_ve_diast_norm,
                    'diam_aortico': diam_aortico,
                    'diam_ae': diam_ae,
                    'diam_atrio_ao_esq': diam_atrio_ao_esq,
                    'vel_pico_va': vel_pico_va,
                    'gp_max_va': gp_max_va,
                    'vel_pico_pulmonar': vel_pico_pulmonar,
                    'gradiente_pico_pulmonar': gradiente_pico_pulmonar,
                    'veloc_pico_mit_onda_e': veloc_pico_mit_onda_e,
                    'veloc_pico_mit_ond_a': veloc_pico_mit_ond_a,
                    'taxa_mitral_e_a': taxa_mitral_e_a,
                    'tapse': tapse,
                    'ventriculo_esq': ventriculo_esq,
                    'massa_ve': massa_ve,
                    'indice_massa_ve': indice_massa_ve,
                    'dia_interno_ve_sist_norm': dia_interno_ve_sist_norm,
                    'espessura_relat_parede': espessura_relat_parede,
                    'tricuspid': tricuspid,
                    'doppler': doppler,
                    'aorta': aorta,
                    'grad_pico_mit_e': grad_pico_mit_e,
                    'grad_pico_mitral_a': grad_pico_mitral_a,
                    'mtp_mitral': mtp_mitral,
                    'regurgitacao_mitral': regurgitacao_mitral,
                    'vel_reg_mit': vel_reg_mit,
                    'grad_reg_mitral': grad_reg_mitral,
                    'dpdt': dpdt,
                    'regurgitacao_tric': regurgitacao_tric,
                    'vel_reg_tric': vel_reg_tric,
                    'grad_reg_tric': grad_reg_tric,
                    'arteria_pulmonar': arteria_pulmonar,
                    'tdi_mitral': tdi_mitral,
                    'pressao_cap_pulmonar': pressao_cap_pulmonar,
                    'aorta_atrio_esq': aorta_atrio_esq,
                    'dia_esfericidade': dia_esfericidade,
                    'diad_diam_ve': diad_diam_ve,
                    'temp_desacel_onda_e_mitral': temp_desacel_onda_e_mitral,
                    'temp_relax_isovol_mitral': temp_relax_isovol_mitral,
                    'temp_relax_isovol_e_mitral': temp_relax_isovol_e_mitral,
                    'notes': notes,
                    'conclusion': conclusion,
                })
                
                # Recalculate reference results with updated data
                weight_num = extract_weight(weight) if weight else None
                if weight_num and weight_num > 0:
                    if is_float(tapse):
                        tapse_fl = float(tapse)
                        tapse_resultado = reference_results("tapse", tapse_fl, weight_num)
                    else:
                        tapse_resultado = ''
                    
                    if is_float(diam_aortico):
                        diam_aortico_result = checkAorta(weight_num, float(diam_aortico))
                    else:
                        diam_aortico_result = ''
                    
                    if is_float(diam_ae):
                        diam_ae_result = checkLeftAtrium(weight_num, float(diam_ae))
                    else:
                        diam_ae_result = ''
                    
                    if is_float(diastole_septo_iv):
                        diastole_septo_iv_result = checkSepto(weight_num, float(diastole_septo_iv))
                    else:
                        diastole_septo_iv_result = ''
                    
                    if is_float(diast_par_post_ve):
                        diast_par_post_ve_result = checkLVWall(weight_num, float(diast_par_post_ve))
                    else:
                        diast_par_post_ve_result = ''
                    
                    if is_float(diast_diad_ve):
                        diast_diad_ve_result = checkLVChamber_d(weight_num, float(diast_diad_ve))
                    else:
                        diast_diad_ve_result = ''
                    
                    if is_float(dia_sist_ve):
                        dia_sist_ve_result = checkLVChamber_s(weight_num, float(dia_sist_ve))
                    else:
                        dia_sist_ve_result = ''
                else:
                    tapse_resultado = ''
                    diam_aortico_result = ''
                    diam_ae_result = ''
                    diastole_septo_iv_result = ''
                    diast_par_post_ve_result = ''
                    diast_diad_ve_result = ''
                    dia_sist_ve_result = ''
                
                st.session_state.extracted_data.update({
                    'tapse_resultado': tapse_resultado,
                    'diam_aortico_result': diam_aortico_result,
                    'diam_ae_result': diam_ae_result,
                    'diastole_septo_iv_result': diastole_septo_iv_result,
                    'diast_par_post_ve_result': diast_par_post_ve_result,
                    'diast_diad_ve_result': diast_diad_ve_result,
                    'dia_sist_ve_result': dia_sist_ve_result,
                })
                
                st.session_state.step = 3
                st.rerun()

    # Step 3: Document Generation
    elif st.session_state.step == 3:
        st.markdown('<h2 class="section-header">📄 Geração do Laudo</h2>', unsafe_allow_html=True)
        
        # Back button
        if st.button("⬅️ Voltar à Confirmação"):
            st.session_state.step = 2
            st.rerun()
        
        extracted_data = st.session_state.extracted_data
        
        st.markdown("---")
        uploaded_images = st.file_uploader(
            "📷 Envie imagens (JPEG/PNG) para incluir no laudo",
            type=["png", "jpg", "jpeg"],
            accept_multiple_files=True,
            help="As imagens serão inseridas em grupos de até 6 por página no final do documento."
        )
        if uploaded_images:
            st.info(f"{len(uploaded_images)} imagem(ns) pronta(s) para inclusão.")
        st.markdown("---")
        
        # Botão para gerar laudo
        if st.button("📄 Gerar Laudo", type="primary"):
            
            with st.spinner("Gerando laudo..."):
                try:
                    # Preparar imagens enviadas pelo usuário
                    images = []
                    temp_image_paths = []
                    if uploaded_images:
                        for idx, img_file in enumerate(uploaded_images):
                            try:
                                ext = os.path.splitext(img_file.name)[1] or ".png"
                                with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_img:
                                    tmp_img.write(img_file.getbuffer())
                                    img_path = tmp_img.name
                                    temp_image_paths.append(img_path)
                                
                                with Image.open(img_path) as img_obj:
                                    width, height = img_obj.size
                                
                                images.append({
                                    "path": img_path,
                                    "page": idx + 1,
                                    "width": width,
                                    "height": height,
                                    "name": img_file.name,
                                })
                            except Exception as e:
                                st.warning(f"Não foi possível usar a imagem {img_file.name}: {e}")
                    
                    if uploaded_images and not images:
                        st.warning("⚠️ Nenhuma imagem válida foi carregada; o laudo será salvo sem imagens.")
                    
                    # Prepare data for document generation
                    data = {
                        '<exam_date>': extracted_data.get('date_of_exam', ''),
                        '<owner_name>': extracted_data.get('owner_name', ''),
                        '<species>': extracted_data.get('species', ''),
                        '<sex>': extracted_data.get('sex', ''),
                        '<age>': extracted_data.get('age', ''),
                        '<weight>': extracted_data.get('weight', ''),
                        '<ficha>': extracted_data.get('identification', ''),
                        '<breed>': extracted_data.get('breed', ''),
                        '<animal_name>': extracted_data.get('animal_name', ''),
                        '<frequencia_cardiaca>': extracted_data.get('frequencia_cardiaca', ''),
                        '<ritmo>': extracted_data.get('ritmo', ''),
                        '<description>': extracted_data.get('description', ''),
                        '<report_date>': extracted_data.get('report_date', ''),
                        '<diastole_septo_IV>': extracted_data.get('diastole_septo_iv', ''),
                        '<diastole_parede_post_VE>': extracted_data.get('diast_par_post_ve', ''),
                        '<diast_diametro_VE>': extracted_data.get('diast_diad_ve', ''),
                        '<diametro_sist_VE>': extracted_data.get('dia_sist_ve', ''),
                        '<encurtamento_fracional_VE>': extracted_data.get('encurtamento_frac_VE', ''),
                        '<fracao_ejecao>': extracted_data.get('frac_ejecao', ''),
                        '<diametro_interno_VE_diast_norm>': extracted_data.get('diad_interno_ve_diast_norm', ''),
                        '<diametro_aortico>': extracted_data.get('diam_aortico', ''),
                        '<diametro_AE>': extracted_data.get('diam_ae', ''),
                        '<diam_atrio_ao_esq>': extracted_data.get('diam_atrio_ao_esq', ''),
                        '<vmax_va>': extracted_data.get('vel_pico_va', ''),
                        '<gp_max_va>': extracted_data.get('gp_max_va', ''),
                        '<vel_pico_pulmonar>': extracted_data.get('vel_pico_pulmonar', ''),
                        '<gradiente_pico_pulmonar>': extracted_data.get('gradiente_pico_pulmonar', ''),
                        '<vel_pico_mitral_e>': extracted_data.get('veloc_pico_mit_onda_e', ''),
                        '<vel_pico_mit_ond_a>': extracted_data.get('veloc_pico_mit_ond_a', ''),
                        '<taxa_mitral_ea>': extracted_data.get('taxa_mitral_e_a', ''),
                        '<temp_desacel_onda_e_mitral>': extracted_data.get('temp_desacel_onda_e_mitral', ''),
                        '<temp_relax_isovol_mitral>': extracted_data.get('temp_relax_isovol_mitral', ''),
                        '<temp_relax_isovol_e_mitral>': extracted_data.get('temp_relax_isovol_e_mitral', ''),
                        '<onda_e_lateral>': extracted_data.get('onda_e_lateral', ''),
                        '<onda_a_lateral>': extracted_data.get('onda_a_lateral', ''),
                        '<razao_e_a_lat>': extracted_data.get('razao_ee_lat', ''),
                        '<tapse>': extracted_data.get('tapse', ''),
                        '<tapse_resultado>': extracted_data.get('tapse_resultado', ''),
                        '<notes>': extracted_data.get('notes', ''),
                        '<conclusion>': extracted_data.get('conclusion', ''),
                        '<diam_aortico_result>': extracted_data.get('diam_aortico_result', ''),
                        '<diastole_septo_iv_result>': extracted_data.get('diastole_septo_iv_result', ''),
                        '<diam_ae_result>': extracted_data.get('diam_ae_result', ''),
                        '<diast_par_post_ve_result>': extracted_data.get('diast_par_post_ve_result', ''),
                        '<diast_diametro_VE_result>': extracted_data.get('diast_diad_ve_result', ''),
                        '<dia_sist_ve_result>': extracted_data.get('dia_sist_ve_result', ''),
                        '<mapse>': extracted_data.get('mapse', ''),
                    }
                    
                    # Template path
                    template_path = os.path.join(os.path.dirname(__file__), 'heartcaresite', 'upload_folder', 'Laudo Eco Modelo P.docx')
                
                    if not os.path.exists(template_path):
                        st.error("❌ Template não encontrado. Verifique se o arquivo 'Laudo Eco Modelo P.docx' está na pasta correta.")
                        st.stop()
                    
                    # Create output filename
                    animal_name = extracted_data.get('animal_name', 'Animal')
                    date_str = extracted_data.get('date_of_exam', '')
                    
                    if date_str:
                        try:
                            date_object = datetime.strptime(date_str, "%d/%m/%Y")
                            formatted_date = date_object.strftime("%d%m%Y")
                        except:
                            formatted_date = datetime.now().strftime("%d%m%Y")
                    else:
                        formatted_date = datetime.now().strftime("%d%m%Y")
                    
                    output_filename = f"Laudo_{animal_name}_{formatted_date}.docx"
                    
                    # Create temporary file for output
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_output:
                        output_path = tmp_output.name
                    
                    # Generate document
                    if create_word_document(data, template_path, output_path, images=images if images else None):
                        # Read the generated file
                        with open(output_path, 'rb') as file:
                            file_content = file.read()
                        
                        # Clean up temporary file
                        os.unlink(output_path)
                        
                        success_msg = "✅ Laudo gerado com sucesso!"
                        if images and len(images) > 0:
                            success_msg += f" ({len(images)} imagem(ns) incluída(s))"
                        
                        st.success(success_msg)
                        
                        st.markdown("""
                        <div class="success-box">
                        <strong>Laudo Ecocardiográfico Gerado!</strong><br>
                        O documento foi criado com todas as informações do exame e referências médicas calculadas.
                        """ + (f"<br>📷 {len(images)} imagem(ns) do PDF foram incluídas." if images and len(images) > 0 else "") + """
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Download button
                        st.download_button(
                            label="📥 Baixar Laudo",
                            data=file_content,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                    
                        # Show summary
                        with st.expander("📋 Resumo do Laudo Gerado", expanded=True):
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                st.write("**Informações do Animal:**")
                                st.write(f"- Nome: {extracted_data.get('animal_name', 'N/A')}")
                                st.write(f"- Espécie: {extracted_data.get('species', 'N/A')}")
                                st.write(f"- Raça: {extracted_data.get('breed', 'N/A')}")
                                st.write(f"- Peso: {extracted_data.get('weight', 'N/A')}")
                                st.write(f"- Data do Exame: {extracted_data.get('date_of_exam', 'N/A')}")
                            
                            with col2:
                                st.write("**Parâmetros Cardíacos:**")
                                if extracted_data.get('frequencia_cardiaca'):
                                    st.write(f"- Frequência Cardíaca: {extracted_data.get('frequencia_cardiaca')} bpm")
                                if extracted_data.get('ritmo'):
                                    st.write(f"- Ritmo: {extracted_data.get('ritmo')}")
                                if extracted_data.get('frac_ejecao'):
                                    st.write(f"- Fração Ejeção: {extracted_data.get('frac_ejecao')}%")
                                if extracted_data.get('tapse'):
                                    st.write(f"- TAPSE: {extracted_data.get('tapse')} mm")
                            
                            with col3:
                                st.write("**Referências Calculadas:**")
                                if extracted_data.get('tapse_resultado'):
                                    st.write(f"- TAPSE: {extracted_data.get('tapse_resultado')}")
                                if extracted_data.get('diam_aortico_result'):
                                    st.write(f"- Diâmetro Aórtico: {extracted_data.get('diam_aortico_result')}")
                                if extracted_data.get('diam_ae_result'):
                                    st.write(f"- Diâmetro AE: {extracted_data.get('diam_ae_result')}")
                                if extracted_data.get('diastole_septo_iv_result'):
                                    st.write(f"- Diástole-Septo IV: {extracted_data.get('diastole_septo_iv_result')}")
                        
                        # Option to start new report
                        if st.button("🔄 Gerar Novo Laudo", type="secondary"):
                            st.session_state.step = 1
                            st.session_state.extracted_data = {}
                            st.session_state.pdf_file = None
                            st.session_state.pdf_bytes = None
                            st.rerun()
                    else:
                        st.error("❌ Erro ao gerar o laudo. Verifique os dados e tente novamente.")
                
                except Exception as e:
                    st.error(f"❌ Erro ao gerar o laudo: {str(e)}")
                    st.exception(e)
                
                finally:
                    for p in temp_image_paths:
                        try:
                            if os.path.exists(p):
                                os.unlink(p)
                        except:
                            pass

    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style="text-align: center; color: #666; font-size: 0.9rem;">
        Vet Heart Care - Sistema de Laudos Ecocardiográficos<br>
        Desenvolvido para veterinários especializados em cardiologia
        </div>
        """,
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()
