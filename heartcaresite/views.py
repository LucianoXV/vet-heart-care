from django.http import HttpResponse,HttpResponseRedirect
from django.shortcuts import render,redirect
from .forms import PDFUploadForm, NameForm, DataConfirmationForm
import pdfplumber
from docx import Document
from datetime import datetime,date
from .models import Owner, Pet, CardiacData
from EcoDataReferences import checkLeftAtrium, checkAorta, checkSepto, checkLVWall, checkLVChamber_d, checkLVChamber_s
import os
import pg8000


owner_name = None
generated_doc_path = None

def index(request):
    return render(request, 'heartcaresite/index.html') 

def get_name(request):
    # if this is a POST request we need to process the form data
    if request.method == "POST":
        # create a form instance and populate it with data from the request:
        form = NameForm(request.POST)
        # check whether it's valid:
        if form.is_valid():
            # process the data in form.cleaned_data as required
            # ...
            
            # redirect to a new URL:
            return HttpResponseRedirect("/thanks/")

    # if a GET (or any other method) we'll create a blank form
    else:
        form = NameForm()

    return render(request, "heartcaresite/name.html", {"form": form})

def upload_pdf(request):
    if request.method == 'POST':
        form = PDFUploadForm(request.POST, request.FILES)
        if form.is_valid():
            pdf_file = form.cleaned_data['pdf_file']
            # Process the PDF file (e.g., extract data)

            # Function to extract table data from a PDF
            table_data = []
            
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    # Extract table data from the page, if any
                    tables = page.extract_tables()
                    for table in tables:
                        table_data.extend(table)
            request.session['table_data'] = table_data     

            return redirect('data_confirmation')
        else:
            # Form is not valid, handle this case (e.g., re-render the form with errors)
            return render(request, 'heartcaresite/upload_pdf.html', {'form': form})
    else:
        form = PDFUploadForm()
    return render(request, 'heartcaresite/upload_pdf.html', {'form': form})

def data_confirmation(request):
    # Retrieve the data object by data_id
    # Display the data to the user with a form for confirmation or modification
    
    table_data = request.session.get('table_data',{})   
    # Initialize variables
    owner_name = animal_name = breed = age = sex = weight = date_of_exam = operator_name = description = report_date =    ""
    diastole_septo_iv = ventriculo_esq = diast_diad_ve = diast_par_post_ve = ""
    dia_sist_ve = frac_ejecao = massa_ve = indice_massa_ve = ""
    diad_interno_ve_diast_norm = dia_interno_ve_sist_norm = ""
    espessura_relat_parede = tricuspid = tapse = doppler = ""
    aorta = vel_pico_va = gp_max_va = veloc_pico_mit_onda_e = veloc_pico_mit_ond_a = ""
    grad_pico_mit_e = grad_pico_mitral_a = mtp_mitral = taxa_mitral_e_a = ""
    temp_desacel_onda_e_mitral = temp_relax_isovol_mitral = temp_relax_isovol_e_mitral = ""
    regurgitacao_mitral = vel_reg_mit = grad_reg_mitral = dpdt = ""
    regurgitacao_tric = vel_reg_tric = grad_reg_tric = arteria_pulmonar = ""
    vel_pico_pulmonar = gradiente_pico_pulmonar = tdi_mitral = onda_e_lateral = onda_a_lateral = ""
    razao_ee_lat = razao_ee_lat = pressao_cap_pulmonar = vel_pico_mit_ond_e = pressao_cap_pulmonar = ""
    aorta_atrio_esq = diam_aortico = diam_ae = dia_esfericidade = diad_diam_ve = ""
    diam_atrio_ao_esq = vmax_va = encurtamento_frac_VE = identification = mapse = ""

    for row in table_data:
        for i in range(0, len(row), 1):
            if row[i] is not None and row[i] != '':
                key = row[i].strip()
                
                value = row[i + 1].strip() if i + 1 < len(row) and row[i + 1] is not None else None
                value2 = row[i + 2].strip() if i + 2 < len(row) and row[i + 2] is not None else None

                ## Uncomment below if you want to check the data that is being read from the pdf file
                print (f"key:{key}")
                print (f"value: {value}")
                print (f'value2: {value2}')

                if key == 'Nome proprietário':
                    owner_name = value
                    owner_name = owner_name.replace("\n", " ")
                    owner_name = owner_name.title()
                elif key == 'Nome do animal':
                    animal_name = value
                    parts = animal_name.split(',')
                    if len(parts)>1:
                        species = parts[1].strip()
                        animal_name = parts[0].strip()
                    else:
                        species = ""
                elif key == 'Raça':
                    breed = value
                elif key == 'Idade':
                    age = value
                elif key == 'Sexo':
                    sex = value
                elif key == 'Peso':
                    weight = value
                elif key == 'Identificação':
                    identification = value
                elif key == 'Data do exame':
                    date_of_exam = value
                elif key == 'Operador':
                    operator_name = value
                elif key == 'Descrição do exame':
                    description = value
                elif key == 'Data do relatório':
                    report_date = value
                elif key == 'Diástole-Septo IV':
                    diastole_septo_iv = value #+ " " + value2
                elif key == 'Ventrículo esq':
                    ventriculo_esq = value
                elif key == 'Diást-diâmetro VE':
                    diast_diad_ve = value #+ " " + value2
                elif key == 'Diástole Parede Post VE':
                    diast_par_post_ve = value #+ " " + value2
                elif key == 'Diâmetro-síst VE':
                    dia_sist_ve = value
                elif key == 'Fração Ejeção':
                    frac_ejecao = value
                elif key == 'MAPSE':
                    mapse = value
                elif key == 'Massa VE':
                    massa_ve = value
                elif key == 'Índice Massa VE':
                    indice_massa_ve = value
                elif key == 'Diâmetro interno VE diást norm':
                    diad_interno_ve_diast_norm = value
                elif key == 'Diâmetro interno VE sist norm':
                    dia_interno_ve_sist_norm = value
                elif key == 'Espessura relativa da parede':
                    espessura_relat_parede = value
                elif key == 'Tricúspid':
                    tricuspid = value
                elif key == 'TAPSE':
                    tapse = value
                elif key == 'Doppler':
                    doppler = value
                elif key == 'Aorta':
                    aorta = value
                elif key == 'Vmáx VA':
                    vel_pico_va = value
                elif key == 'GP máx VA':
                    gp_max_va = value
                elif key == 'Vel Pico Mitral Onda E':
                    veloc_pico_mit_onda_e = value
                elif key == 'Vel Pico Mit Ond A':
                    veloc_pico_mit_ond_a = value
                elif key == 'Grad Pico Mit (E)':
                    grad_pico_mit_e = value
                elif key == 'Gtad Pico Mitral (A)':
                    grad_pico_mitral_a = value
                elif key == 'MTP Mitral':
                    mtp_mitral = value
                elif key == 'Taxa Mitral E/A':
                    taxa_mitral_e_a = value
                elif key == 'Temp Desacel Onda e Mitral':
                    temp_desacel_onda_e_mitral = value
                elif key == 'Temp Relax Isovol Mitral':
                    temp_relax_isovol_mitral = value
                elif key == 'Temp Relax Isovol E Mitral':
                    temp_relax_isovol_e_mitral = value
                elif key == 'Regurgitação Mitral':
                    regurgitacao_mitral = value
                elif key == 'Velocidade Reg Mit':
                    vel_reg_mit = value
                elif key == 'Gradiente Reg Mitral':
                    grad_reg_mitral = value
                elif key == 'dP/dt':
                    dpdt = value
                elif key == 'Regurgitação Tric':
                    regurgitacao_tric = value
                elif key == 'Vel Reg Tric':
                    vel_reg_tric = value
                elif key == 'Grad Reg Tric':
                    grad_reg_tric = value
                elif key == 'Artéria Pulmonar':
                    arteria_pulmonar = value
                elif key == 'Vel pico Pulmonar':
                    vel_pico_pulmonar = value
                elif key == 'Gradiente pico Pulmonar':
                    gradiente_pico_pulmonar = value
                elif key == 'TDI Mitral':
                    tdi_mitral = value
                elif key == 'Onda E\' Lateral':
                    onda_e_lateral = value
                elif key == 'Onda A\' Lateral':
                    onda_a_lateral = value
                elif key == 'Razão E’/A\' Lat':
                    razao_ee_lat = value
                elif key == 'Pressão capilar pulmonar':
                    pressao_cap_pulmonar = value
                elif key == 'Pressão capilar pulmonar':
                    pressao_cap_pulmonar = value
                elif key == 'Aorta/átrio esq':
                    aorta_atrio_esq = value
                elif key == 'Diâmetro Aórtico':
                    diam_aortico = value
                elif key == 'Diâmetro AE':
                    diam_ae = value
                elif key == 'Diâm Átrio/Ao Esq':
                    dia_esfericidade = value
                elif key == 'Índice de esfericidade':
                    diad_diam_ve = value
                elif key == 'Encurtamento Fracional VE':
                    encurtamento_frac_VE = value
                elif key == 'Diâm Átrio/Ao Esq':
                    diam_atrio_ao_esq = value
    
    if len(species) != 0:
        if species == 'FEL':
            species = 'Felina'
        elif species == 'CAN':
            species = 'Canina'
        if sex == 'F':
            sex = 'Fêmea'
        elif sex == 'M':
            sex = 'Macho'

    if len(weight) != 0:                
        weight_num = extract_weight(weight)
        if is_float(tapse):
            tapse_fl = float(tapse)
            tapse_resultado = reference_results("tapse",tapse_fl, weight_num)
        else:
            tapse_resultado = ''
        diam_aortico_result = checkAorta(weight_num,float(diam_aortico))
        diam_ae_result = checkLeftAtrium(weight_num,float(diam_ae))
        diastole_septo_iv_result = checkSepto(weight_num,float(diastole_septo_iv))

        diast_par_post_ve_result = checkLVWall(weight_num,float(diast_par_post_ve))
        diast_diad_ve_result = checkLVChamber_d(weight_num,float(diast_diad_ve))
        dia_sist_ve_result = checkLVChamber_s(weight_num,float(dia_sist_ve))
    else:
        weight = ""
        tapse_resultado = ''
        diam_aortico_result = ''
        diam_ae_result = ''
        diastole_septo_iv_result = ''
        diast_par_post_ve_result = ''
        diast_diad_ve_result = ''
        dia_sist_ve_result = ''

    similar_owners = Owner.objects.filter(name__icontains=owner_name)
    
    form = DataConfirmationForm(initial={'owner_name': owner_name,'animal_name':animal_name,'age':age,'species':species, 'breed':breed, 'sex':sex, 'weight':weight, 'date_of_exam':date_of_exam, 'tapse':tapse,'tapse_resultado':tapse_resultado, 'record_number':identification})

    if request.method == 'POST':
        # Process the form submission
        form = DataConfirmationForm(request.POST)
        if form.is_valid():
            # Save the modified data to the database
            # Your save logic goes here
            # ...
            BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            document_path = os.path.join(BASE_DIR, 'heartcaresite/upload_folder', 'Laudo Eco Modelo P.docx')
            
            doc = Document(document_path)
            data = {
                '<exam_date>': request.POST.get('date_of_exam'),
                '<owner_name>': request.POST.get('owner_name'),
                '<species>': request.POST.get('species'),
                '<sex>': request.POST.get('sex'),
                '<age>': request.POST.get('age'),
                '<weight>': request.POST.get('weight'),
                '<ficha>': request.POST.get('record_number'),
                '<breed>': request.POST.get('breed'),
                '<animal_name>': request.POST.get('animal_name'),
                '<diastole_septo_IV>': diastole_septo_iv,
                '<diastole_parede_post_VE>': diast_par_post_ve,
                '<diast_diametro_VE>': diast_diad_ve,
                '<diametro_sist_VE>': dia_sist_ve,
                '<encurtamento_fracional_VE>': encurtamento_frac_VE,
                '<fracao_ejecao>': frac_ejecao,
                '<diametro_interno_VE_diast_norm>': diad_interno_ve_diast_norm,
                '<diametro_aortico>': diam_aortico,
                '<diametro_AE>': diam_ae,
                '<diam_atrio_ao_esq>': diam_atrio_ao_esq,
                '<vmax_va>': vel_pico_va,
                '<gp_max_va>': gp_max_va,
                '<vel_pico_pulmonar>': vel_pico_pulmonar,
                '<gradiente_pico_pulmonar>': gradiente_pico_pulmonar,
                '<vel_pico_mitral_e>': veloc_pico_mit_onda_e,
                '<vel_pico_mit_ond_a>': veloc_pico_mit_ond_a,
                '<taxa_mitral_ea>': taxa_mitral_e_a,
                '<temp_desacel_onda_e_mitral>': temp_desacel_onda_e_mitral,
                '<temp_relax_isovol_mitral>': temp_relax_isovol_mitral,
                '<temp_relax_isovol_e_mitral>': temp_relax_isovol_e_mitral,
                '<onda_e_lateral>': onda_e_lateral,
                '<onda_a_lateral>': onda_a_lateral,
                '<razao_e_a_lat>': razao_ee_lat,
                '<tapse>': tapse,
                '<tapse_resultado>': tapse_resultado,
                '<notes>': request.POST.get('notes'),
                '<conclusion>': request.POST.get('conclusion'),
                '<diam_aortico_result>': diam_aortico_result,
                '<diastole_septo_iv_result>': diastole_septo_iv_result,
                '<diam_ae_result>': diam_ae_result,
                '<diast_par_post_ve_result>': diast_par_post_ve_result,
                '<diast_diametro_VE_result>': diast_diad_ve_result,
                '<dia_sist_ve_result>': dia_sist_ve_result,
                '<mapse>': mapse
            }
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraphs in cell.paragraphs:
                            for run in paragraphs.runs:
                                run.text = replace_text(run.text,data)

            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    if key in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(key,value)
            
            date_object = datetime.strptime(request.POST.get('date_of_exam'),"%d/%m/%Y")
            formatted_date = date_object.strftime("%d%m%Y")

            doc.save('heartcaresite/laudos_gerados/Laudo_'+animal_name+'_'+formatted_date+'.docx')
            # Redirect or render a success page
            generated_doc_path = BASE_DIR + '/heartcaresite/laudos_gerados/Laudo_'+animal_name+'_'+formatted_date+'.docx'
            request.session['file_path'] = generated_doc_path
            
            owner_name = request.POST.get('owner_name')
            animal_name = request.POST.get('animal_name')
            email = request.POST.get('email')
            phone = request.POST.get('phone')
            address = request.POST.get('address')
            weight=request.POST.get('weight')
            notes = request.POST.get('notes')
            conclusion = request.POST.get('conclusion')
            exam_date = date_object.strftime('%Y-%m-%d')
            today_date = datetime.strptime(str(date.today()), '%Y-%m-%d').date()
            
            # Get or create Owner
            #owner, created_owner = Owner.objects.get_or_create(name=owner_name,email=email, phone=phone, address=address, rec_date=today_date)

            # Check if the owner with the given name exists
            existing_owner = Owner.objects.filter(name=owner_name).first()

            if existing_owner:
                # Object exists, do something with it
                # You can access existing_owner directly
                #owner = existing_owner.id
                owner = Owner.objects.get(id=existing_owner.id)
            else:
                # Object doesn't exist, create a new one with additional attributes
                new_owner = Owner.objects.create(name=owner_name,email=email, phone=phone, address=address, rec_date=today_date)
                owner = new_owner

            existing_pet = Pet.objects.filter(name=animal_name).first()
            if existing_pet:
                pet = existing_pet
            else:
                new_pet = Pet.objects.create(name=animal_name, owner=owner, breed=breed, gender=".", weight=weight,rec_date=today_date)
                pet = new_pet

            # Get or create Pet associated with the Owner
            #pet, created_pet = Pet.objects.get_or_create(name=animal_name, owner=owner, breed=breed, gender=".", weight_kg=request.POST.get('weight'),rec_date=today_date)
            
            # Create CardiacData
            cardiac_data = CardiacData.objects.create(
                pet=pet,
                exam_date=exam_date,
                rec_date=today_date,
                notes=notes,
                conclusion=conclusion,
                weight=weight,
                diastole_septo_iv=string_to_float(diastole_septo_iv),
                diast_par_post_ve=string_to_float(diast_par_post_ve),
                diast_diad_ve = string_to_float(diast_diad_ve),
                dia_sist_ve = string_to_float(dia_sist_ve), 
                encurtamento_frac_VE = string_to_float(encurtamento_frac_VE),
                frac_ejecao = string_to_float(frac_ejecao),
                diad_interno_ve_diast_norm = string_to_float(diad_interno_ve_diast_norm),
                diam_aortico = string_to_float(diam_aortico),
                diam_ae = string_to_float(diam_ae),
                diam_atrio_ao_esq = string_to_float(diam_atrio_ao_esq),
                vel_pico_va = string_to_float(vel_pico_va),

                gp_max_va = string_to_float(gp_max_va),
                vel_pico_pulmonar = string_to_float(vel_pico_pulmonar),
                gradiente_pico_pulmonar = string_to_float(gradiente_pico_pulmonar),
                veloc_pico_mit_onda_e = string_to_float(veloc_pico_mit_onda_e),
                veloc_pico_mit_ond_a = string_to_float(veloc_pico_mit_ond_a),
                taxa_mitral_e_a = string_to_float(taxa_mitral_e_a),
                temp_desacel_onda_e_mitral = string_to_float(temp_desacel_onda_e_mitral),
                temp_relax_isovol_e_mitral = string_to_float(temp_relax_isovol_e_mitral),
                onda_e_lateral = string_to_float(onda_e_lateral),
                onda_a_lateral = string_to_float(onda_a_lateral),
                razao_ee_lat= string_to_float(razao_ee_lat),
                tapse = string_to_float(tapse),
                tapse_resultado = tapse_resultado
                
                #exam_date= datetime.strptime(request.POST.get('date_of_exam'),'%Y-%m-%d')
                # Other fields related to cardiac data
            )
            
            return redirect('view_save_doc')
            
        
    return render(request, 'heartcaresite/data_confirmation.html', {'form': form})

def replace_text(text,data):
    for key, value in data.items():
        text = text.replace(key,value)
    return text

def reference_results(echo_label, echo_value, echo_weight):
    conn = pg8000.connect(user="postgres", password="0375lxv", host="localhost", database="VetHeartCare")
    cursor = conn.cursor()
    select_query = """
        SELECT max(min)
        FROM heartcaresite_tapse
        WHERE weight <= %s
    """
    condition_value = echo_weight
    # Execute the SELECT query with the provided values
    cursor.execute(select_query, (condition_value,))
    result_min = cursor.fetchone()  # Use fetchall() if you expect multiple rows

    select_query = """
        SELECT min(max)
        FROM heartcaresite_tapse
        WHERE weight >= %s
    """
    condition_value = echo_weight
    # Execute the SELECT query with the provided values
    
    cursor.execute(select_query, (condition_value,))
    result_max = cursor.fetchone()  # Use fetchall() if you expect multiple rows

    print (f'conditional value: {condition_value}')
    print (f'result_max: {result_max}')

    # Close the cursor and connection
    cursor.close()
    conn.close()

    # Now you can use the 'result' variable to access the values from the query
    if result_min:
        column1_value = result_min
        column2_value = result_max
    else:
        print("No results found.")

    if echo_value > float(result_min[0]) and echo_value < float(result_max[0]):
        result = "Normal"
    elif echo_value < float(result_min[0]):
        result = "Diminuido"
    elif echo_value > float(result_max[0]):
        result = "Aumentado"
    return result

def extract_weight(input_string):
    # Find the index of the string "kg"
    kg_index = input_string.find("kg")

    # Extract the substring before "kg" (excluding leading and trailing whitespaces)
    weight_str = input_string[:kg_index].strip()

    try:
        # Convert the extracted substring to an integer
        weight_int = int(weight_str)
        return weight_int
    except ValueError:
        # Handle the case where the conversion to integer fails
        print(f"Error: Unable to convert '{weight_str}' to an integer.")
        return None

def voltar_action(request):
    return redirect('upload_pdf')

def view_save_doc (request):
    print('got here')
    if request.method == 'GET':
        # Process the form submission
        form = DataConfirmationForm(request.GET)
        if form.is_valid():
            file_path = request.session.get('file_path')
            if file_path:  # Check if file_path is not None
                if os.path.exists(file_path):
                    with open(file_path, 'rb') as file:
                        response = HttpResponse(file.read(), content_type='application/octet-stream')
                        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
                        #return response
                        return redirect('success_page')
                else:
                    return HttpResponse("File not found")
            else:
                return HttpResponse("File path is None")
        print('got here3')
        

def success_page(request):
    return render(request, 'heartcaresite/docx_successpage.html')           


def is_float(value):
    try:
        float_value = float(value)
        return True  # If conversion to float succeeds
    except ValueError:
        return False  #  If conversion to float fails

def string_to_float(string_value):
    print(string_value)
    if string_value.strip():  # Check if the string is not empty after stripping whitespace
        try:
            return float(string_value)
        except ValueError:
            # Handle the case where the string cannot be converted to a float
            return None  # Or raise an exception, log an error, etc.
    else:
        # Handle the case where the string is empty
        return None  # Or any default value you want to assign